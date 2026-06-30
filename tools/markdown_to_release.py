from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import os
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PYTHON_RUNTIME_ROOT = Path("/Users/xiaochen/.cache/codex-runtimes/codex-primary-runtime/dependencies")
SOFFICE = PYTHON_RUNTIME_ROOT / "bin/soffice"


@dataclass(frozen=True)
class ReleaseConfig:
    lang: str
    title: str
    subtitle: str
    source_dir: Path
    docx_path: Path
    pdf_path: Path
    font: str


BOOK_FILES = [
    "00-preface.md",
    "01-classical-computer-engineering.md",
    "02-llm-compute-engine.md",
    "03-agent-orchestrator.md",
    "99-future-chapters.md",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_font(paragraph, font_name: str) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def configure_styles(doc: Document, font_name: str) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 18, "1F4D78", 18, 8),
        ("Heading 2", 14, "2E74B5", 14, 6),
        ("Heading 3", 12, "1F4D78", 10, 4),
    ]:
        style = styles[name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Number", "List Bullet"):
        style = styles[list_style]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_title(doc: Document, config: ReleaseConfig) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(config.title)
    run.font.name = config.font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), config.font)
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(config.subtitle)
    run.font.name = config.font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), config.font)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)


def iter_book_lines(source_dir: Path) -> list[str]:
    lines: list[str] = []
    for filename in BOOK_FILES:
        text = (source_dir / filename).read_text(encoding="utf-8")
        if filename != BOOK_FILES[0]:
            lines.extend(["", "\\pagebreak", ""])
        lines.extend(text.splitlines())
    return lines


def split_table_row(line: str) -> list[str]:
    raw = line.strip().strip("|")
    cells = re.split(r"(?<!\\)\|", raw)
    return [cell.replace(r"\|", "|").strip() for cell in cells]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def add_markdown_table(doc: Document, rows: list[list[str]], font_name: str) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, 9360)
    set_cell_margins(table)
    col_width = max(900, int(9360 / width))
    for row_idx, row in enumerate(rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, col_width)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            run.font.size = Pt(9 if width >= 3 else 10)
            if row_idx == 0:
                run.bold = True
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], font_name: str) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)


def add_image(doc: Document, image_path: Path, alt: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    if alt:
        caption = doc.add_paragraph(alt)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)


def add_text_paragraph(doc: Document, line: str, font_name: str) -> None:
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(line[2:])
        run.italic = True
        run.font.color.rgb = RGBColor(31, 77, 120)
    elif re.match(r"^\d+\. ", line):
        p = doc.add_paragraph(line, style="List Number")
    elif line.startswith("- "):
        p = doc.add_paragraph(line[2:], style="List Bullet")
    elif line.startswith("# "):
        p = doc.add_paragraph(line[2:], style="Heading 1")
    elif line.startswith("## "):
        p = doc.add_paragraph(line[3:], style="Heading 2")
    elif line.startswith("### "):
        p = doc.add_paragraph(line[4:], style="Heading 3")
    elif line.startswith("*") and line.endswith("*") and len(line) > 2:
        p = doc.add_paragraph()
        run = p.add_run(line.strip("*"))
        run.italic = True
        run.font.color.rgb = RGBColor(85, 85, 85)
    else:
        p = doc.add_paragraph(line)
    set_paragraph_font(p, font_name)


def build_docx(config: ReleaseConfig) -> None:
    config.docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc, config.font)

    lines = iter_book_lines(config.source_dir)
    # The first markdown title/subtitle are already represented by add_title.
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    if lines and lines[0] == "":
        lines = lines[1:]
    subtitle_from_md = None
    if lines and lines[0] and not lines[0].startswith("#"):
        subtitle_from_md = lines[0]
        lines = lines[1:]
    add_title(doc, config)
    if subtitle_from_md and subtitle_from_md != config.subtitle:
        doc.add_paragraph(subtitle_from_md)

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if not line.strip():
            idx += 1
            continue
        if line == "\\pagebreak":
            doc.add_section(WD_SECTION.NEW_PAGE)
            idx += 1
            continue
        if line.startswith("```"):
            code_lines: list[str] = []
            idx += 1
            while idx < len(lines) and not lines[idx].startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code_block(doc, code_lines, config.font)
            idx += 1
            continue
        if line.startswith("|") and idx + 1 < len(lines) and is_table_separator(lines[idx + 1]):
            table_rows = [split_table_row(line)]
            idx += 2
            while idx < len(lines) and lines[idx].startswith("|"):
                table_rows.append(split_table_row(lines[idx]))
                idx += 1
            add_markdown_table(doc, table_rows, config.font)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            alt, rel_path = image_match.groups()
            add_image(doc, config.source_dir / rel_path, alt)
            idx += 1
            continue
        add_text_paragraph(doc, line, config.font)
        idx += 1

    doc.save(config.docx_path)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        lo_profile = tmp_path / "lo-profile"
        cache_home = tmp_path / "cache"
        lo_profile.mkdir(parents=True, exist_ok=True)
        cache_home.mkdir(parents=True, exist_ok=True)
        env = os.environ.copy()
        env["HOME"] = str(tmp_path)
        env["XDG_CACHE_HOME"] = str(cache_home)
        subprocess.run(
            [
                str(SOFFICE),
                f"-env:UserInstallation=file://{lo_profile}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_path),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=env,
        )
        generated = tmp_path / f"{docx_path.stem}.pdf"
        shutil.copyfile(generated, pdf_path)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    version = "v0.3"
    configs = [
        ReleaseConfig(
            lang="zh",
            title="AI Agent 系统设计",
            subtitle="从经典计算机工程到现代智能体架构 - 第一部分：前三章完整版",
            source_dir=root / "docs/zh",
            docx_path=root / f"releases/docx/AI_Agent_System_Design_CN_Part1_Ch1-3_{version}.docx",
            pdf_path=root / f"releases/pdf/AI_Agent_System_Design_CN_Part1_Ch1-3_{version}.pdf",
            font="Arial",
        ),
        ReleaseConfig(
            lang="en",
            title="AI Agent System Design",
            subtitle="From Classical Computer Engineering to Modern Agent Architectures - Part I: Chapters 1-3",
            source_dir=root / "docs/en",
            docx_path=root / f"releases/docx/AI_Agent_System_Design_EN_Part1_Ch1-3_{version}.docx",
            pdf_path=root / f"releases/pdf/AI_Agent_System_Design_EN_Part1_Ch1-3_{version}.pdf",
            font="Arial",
        ),
    ]
    for config in configs:
        build_docx(config)
        convert_docx_to_pdf(config.docx_path, config.pdf_path)
        print(config.docx_path)
        print(config.pdf_path)


if __name__ == "__main__":
    main()
